"""Hand-labeled ground truth across all four formats -> precision/recall/F1.

Builds one small, fully-labeled fixture per format (JSON/DOCX/TXT scored by
exact node path since those are deterministic; PDF scored by value substring
since pdfplumber's line-clustering order isn't guaranteed stable), runs each
through the same detection core `redact.py` uses, and writes eval_report.md.
"""

import json
import os
import tempfile

import docx
import fitz

import parsers
from redact import redact_document


def _json_fixture(tmpdir):
    path = os.path.join(tmpdir, "ticket.json")
    data = {
        "ticket_id": "TCKT-88213",
        "customer_name": "Rohit Kushal Hegde",
        "contact": {"email": "rohit.hegde@examplecorp.com", "phone": "+91 9876543210"},
        "notes": "Customer Shri Rajesh Kumar called about his order.",
        "billing_address": "12 MG Road, Pune 411001",
        "din": "00135070",
        "plan_id": "PLAN-42",
    }
    with open(path, "w") as f:
        json.dump(data, f, indent=2)
    truth = [
        ("ticket_id", None),
        ("customer_name", "name"),
        ("contact.email", "email"),
        ("contact.phone", "phone"),
        ("notes", "name"),
        ("billing_address", "address"),
        ("din", None),
        ("plan_id", None),
    ]
    return path, truth


def _txt_fixture(tmpdir):
    path = os.path.join(tmpdir, "sample.txt")
    lines = [
        "Customer Name: Priya Nair",
        "Email: priya.nair@examplecorp.com",
        "Phone - 9876501234",
        "Order #: 5567123",
        "Ticket No: TCKT-99123",
        "Address: 18 FC Road, Pune 411005",
        "DOB: 15-08-1985",
        "",
        "The applicant, Smt Anjali Deshpande, D/o Vinod Deshpande, resident of Kothrud, submitted the form.",
        "Registered Office at 11/3, Village Birdewadi, Pune - 410501.",
    ]
    with open(path, "w") as f:
        f.write("\n".join(lines) + "\n")
    truth = [
        ("line[0]", "name"),
        ("line[1]", "email"),
        ("line[2]", "phone"),
        ("line[3]", None),
        ("line[4]", None),
        ("line[5]", "address"),
        ("line[6]", "dob"),
        ("line[8]", "name"),
        ("line[9]", None),
    ]
    return path, truth


def _docx_fixture(tmpdir):
    path = os.path.join(tmpdir, "board.docx")
    d = docx.Document()
    d.add_paragraph("Telephone: 022-68052182 Email: ipocmg@icicibank.com")
    d.add_paragraph("Registered Office at 11/3, Village Birdewadi, Pune - 410501.")
    t = d.add_table(rows=1, cols=4)
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Name", "Designation", "DIN", "Address"
    rows = [
        ("Kushal Subbayya Hegde", "Chairman", "00135070", "S. no. 245/104, Pushpakamal, Pune"),
        ("XYZ Family Trust", "Promoter Entity", "N/A", "Registered Office, Pune"),
    ]
    for name, desig, din, addr in rows:
        row = t.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = name, desig, din, addr
    d.save(path)
    # embedded pseudo-fields ("Telephone: ... Email: ...") get folded back
    # into their parent paragraph's own path at detection time, so both
    # expectations key off the same "body.paragraph[0]" path
    truth = [
        ("body.paragraph[0]", "phone"),
        ("body.paragraph[0]", "email"),
        ("body.paragraph[1]", None),
        ("table[0].row[1].cell[0]", "name"),
        ("table[0].row[1].cell[3]", "address"),
        ("table[0].row[1].cell[2]", None),
        ("table[0].row[2].cell[0]", None),
        ("table[0].row[2].cell[3]", None),
    ]
    return path, truth


def _pdf_fixture(tmpdir):
    path = os.path.join(tmpdir, "sample.pdf")
    fitz_doc = fitz.open()
    page = fitz_doc.new_page()
    lines = [
        "Name: Priya Nair",
        "Email: priya.nair@examplecorp.com",
        "Phone: 9876501234",
        "Order Number: 5567123",
    ]
    y = 72
    for line in lines:
        page.insert_text((72, y), line, fontsize=11)
        y += 20
    fitz_doc.save(path)
    fitz_doc.close()
    truth_substrings = [
        ("Priya Nair", "name"),
        ("priya.nair@examplecorp.com", "email"),
        ("9876501234", "phone"),
        ("5567123", None),
    ]
    return path, truth_substrings


def _score_by_path(root, truth):
    leaf_spans, _ = redact_document(root)
    found = {}
    for node, spans in leaf_spans:
        found.setdefault(node.path, set()).update(sp.pii_type for sp in spans)

    tp = fn = fp = tn = 0
    rows = []
    for path, expected in truth:
        detected_types = found.get(path, set())
        if expected is None:
            if detected_types:
                fp += 1
                rows.append((path, expected, detected_types, "FP"))
            else:
                tn += 1
                rows.append((path, expected, detected_types, "TN"))
        elif expected in detected_types:
            tp += 1
            rows.append((path, expected, detected_types, "TP"))
        else:
            fn += 1
            rows.append((path, expected, detected_types, "FN"))
    return tp, fn, fp, tn, rows


def _score_by_substring(root, truth_substrings):
    leaf_spans, _ = redact_document(root)
    all_spans = [(sp.text, sp.pii_type) for _, spans in leaf_spans for sp in spans]

    tp = fn = fp = tn = 0
    rows = []
    for needle, expected in truth_substrings:
        matched_types = {t for text, t in all_spans if needle in text or text in needle}
        if expected is None:
            if matched_types:
                fp += 1
                rows.append((needle, expected, matched_types, "FP"))
            else:
                tn += 1
                rows.append((needle, expected, matched_types, "TN"))
        elif expected in matched_types:
            tp += 1
            rows.append((needle, expected, matched_types, "TP"))
        else:
            fn += 1
            rows.append((needle, expected, matched_types, "FN"))
    return tp, fn, fp, tn, rows


def main():
    fixtures = []
    with tempfile.TemporaryDirectory() as tmpdir:
        path, truth = _json_fixture(tmpdir)
        root, _ = parsers.parse_json(path)
        fixtures.append(("json", path, _score_by_path(root, truth)))

        path, truth = _txt_fixture(tmpdir)
        root, _ = parsers.parse_txt(path)
        fixtures.append(("txt", path, _score_by_path(root, truth)))

        path, truth = _docx_fixture(tmpdir)
        root, _ = parsers.parse_docx(path)
        fixtures.append(("docx", path, _score_by_path(root, truth)))

        path, truth = _pdf_fixture(tmpdir)
        root, _ = parsers.parse_pdf(path)
        fixtures.append(("pdf", path, _score_by_substring(root, truth)))

    total_tp = sum(f[2][0] for f in fixtures)
    total_fn = sum(f[2][1] for f in fixtures)
    total_fp = sum(f[2][2] for f in fixtures)
    total_tn = sum(f[2][3] for f in fixtures)

    precision = total_tp / (total_tp + total_fp) if (total_tp + total_fp) else float("nan")
    recall = total_tp / (total_tp + total_fn) if (total_tp + total_fn) else float("nan")
    f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) else float("nan")

    lines = [
        "# Evaluation Report\n\n",
        "Hand-labeled fixtures across all four formats (JSON, DOCX, PDF, TXT), "
        "run through the same detection core used by `redact.py`.\n\n",
        "## Overall\n\n",
        "| Metric | Value |\n|---|---|\n",
        f"| True Positives | {total_tp} |\n",
        f"| False Negatives | {total_fn} |\n",
        f"| False Positives | {total_fp} |\n",
        f"| True Negatives (correctly left alone) | {total_tn} |\n",
        f"| Precision | {precision:.2f} |\n",
        f"| Recall | {recall:.2f} |\n",
        f"| F1 | {f1:.2f} |\n",
    ]

    for fmt, path, (tp, fn, fp, tn, rows) in fixtures:
        lines.append(f"\n## {fmt.upper()} fixture (`{os.path.basename(path)}`)\n\n")
        lines.append("| Field / value | Expected | Detected | Result |\n|---|---|---|---|\n")
        for field, expected, detected, result in rows:
            lines.append(
                f"| `{field}` | {expected or '_(none)_'} | "
                f"{', '.join(detected) or '_(none)_'} | {result} |\n"
            )

    lines.append("\n## Notes on observed failure modes\n\n")
    lines.append(
        "- Name detection intentionally avoids bare spaCy PERSON hits (weak recall on "
        "Indian names) in favor of context+shape and honorific/kinship patterns; recall "
        "on names with neither an explicit label nor an honorific/kinship marker nearby "
        "is a known gap.\n"
        "- Free-flowing prose addresses with no label or table-column context are out of "
        "scope for the balanced tier — only labeled/columnar address fields are redacted.\n"
        "- The PDF fixture is scored by value substring rather than exact node path, "
        "since pdfplumber's line-clustering order isn't guaranteed stable across inputs.\n"
    )

    with open("eval_report.md", "w", encoding="utf-8") as f:
        f.writelines(lines)

    print(f"Precision: {precision:.2f}  Recall: {recall:.2f}  F1: {f1:.2f}")
    print("Wrote eval_report.md")


if __name__ == "__main__":
    main()
