"""Format-specific parsers that all normalize into the same Node tree.

Every parser returns (root: Node, handle) where `handle` is whatever the
matching write_* function needs to commit redactions back into the original
document. Detection code (detectors.py) never looks at `handle` — it only
ever reads Node.label, Node.path and Node.value.
"""

import json
import re
from dataclasses import dataclass, field
from typing import Any, Optional

import docx
from docx.table import Table as DocxTable

import pdfplumber
import fitz  # PyMuPDF


# ---------------------------------------------------------------------------
# Shared node schema
# ---------------------------------------------------------------------------

@dataclass
class Node:
    id: str
    format: str                 # "json" | "docx" | "pdf" | "txt"
    node_type: str              # see per-format enumerations in the plan
    label: Optional[str]        # the context signal: key / column header / pseudo-field label
    path: str                   # human-readable breadcrumb, used in logs + eval
    value: Optional[str]        # leaf text; None for containers
    children: list = field(default_factory=list)
    source_ref: Any = None      # opaque handle for the writer to commit edits
    meta: dict = field(default_factory=dict)

    def is_leaf(self):
        return self.value is not None

    def walk(self):
        yield self
        for c in self.children:
            yield from c.walk()


# ---------------------------------------------------------------------------
# Format detection (extension vs content sniff must agree)
# ---------------------------------------------------------------------------

class FormatMismatch(Exception):
    def __init__(self, path, ext_guess, sniff_guess):
        self.path, self.ext_guess, self.sniff_guess = path, ext_guess, sniff_guess
        super().__init__(
            f"{path}: extension suggests '{ext_guess}' but content looks like '{sniff_guess}'"
        )


def _guess_from_extension(path):
    ext = path.lower().rsplit(".", 1)[-1] if "." in path else ""
    return {"json": "json", "docx": "docx", "pdf": "pdf", "txt": "txt"}.get(ext, "unknown")


def _sniff_content(path):
    with open(path, "rb") as f:
        head = f.read(8)
    if head.startswith(b"%PDF-"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):
        return "docx"  # docx is a zip; good enough without inspecting [Content_Types].xml
    try:
        with open(path, "r", encoding="utf-8") as f:
            json.load(f)
        return "json"
    except Exception:
        return "txt"


def detect_format(path):
    ext_guess = _guess_from_extension(path)
    sniff_guess = _sniff_content(path)
    if ext_guess != "unknown" and ext_guess != sniff_guess:
        raise FormatMismatch(path, ext_guess, sniff_guess)
    return sniff_guess if ext_guess == "unknown" else ext_guess


# ---------------------------------------------------------------------------
# Shared pseudo-field extraction: "Label: value" segments embedded in prose
# ---------------------------------------------------------------------------

_LABEL_TOKEN_RE = re.compile(
    r"\b([A-Z][A-Za-z]{1,25}(?:[ /][A-Z][A-Za-z]{1,25}){0,3})\s*:\s*"
)


def extract_pseudo_fields(text):
    """Find embedded 'Label: value' spans within a block of prose.

    Returns a list of (label, val_start, val_end, val_text). Used for DOCX
    paragraphs / PDF lines where multiple fields are concatenated on one
    line, e.g. "Telephone: 022-680... Email: x@y.com".
    """
    matches = list(_LABEL_TOKEN_RE.finditer(text))
    fields = []
    for i, m in enumerate(matches):
        label = m.group(1).strip()
        val_start = m.end()
        val_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        raw = text[val_start:val_end]
        stripped = raw.strip()
        if not stripped:
            continue
        offset = raw.find(stripped)
        real_start = val_start + offset
        real_end = real_start + len(stripped)
        fields.append((label, real_start, real_end, stripped))
    return fields


_TXT_KV_RE = re.compile(r"^\s*([A-Za-z][\w /&]{1,40}?)\s*:\s*(.+?)\s*$")
_TXT_DASH_RE = re.compile(r"^\s*([A-Za-z][\w /&]{1,40}?)\s*-\s*(.+?)\s*$")


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------

def parse_json(path):
    with open(path, "r", encoding="utf-8") as f:
        raw = json.load(f)

    counter = [0]

    def walk(obj, label, path_str, parent, key):
        counter[0] += 1
        nid = f"json:{counter[0]}"
        if isinstance(obj, dict):
            node = Node(nid, "json", "object", label, path_str, None,
                        source_ref=(parent, key))
            for k, v in obj.items():
                child_path = f"{path_str}.{k}" if path_str else k
                node.children.append(walk(v, k, child_path, obj, k))
            return node
        if isinstance(obj, list):
            node = Node(nid, "json", "array", label, path_str, None,
                        source_ref=(parent, key))
            for i, v in enumerate(obj):
                node.children.append(walk(v, label, f"{path_str}[{i}]", obj, i))
            return node
        text = "" if obj is None else str(obj)
        return Node(nid, "json", "field", label, path_str, text,
                    source_ref=(parent, key),
                    meta={"atomic": True, "orig_type": type(obj).__name__})

    root = walk(raw, None, "", None, None)
    return root, raw


def write_json(root, raw, output_path):
    for node in root.walk():
        spans = node.meta.get("redactions")
        if not spans or node.value is None:
            continue
        new_value = _splice(node.value, spans)
        parent, key = node.source_ref
        if parent is not None:
            parent[key] = new_value
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(raw, f, indent=2, ensure_ascii=False)


def _splice(text, spans):
    """Apply [(start, end, replacement), ...] to text, highest offset first."""
    out = text
    for start, end, repl in sorted(spans, key=lambda s: s[0], reverse=True):
        out = out[:start] + repl + out[end:]
    return out


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

_SIGNATURE_ZONE_RE = re.compile(
    # deliberately excludes the bare words "signature"/"signed" — they match
    # proper nouns like "Signature Building" (a real address in the RHP
    # fixture set), which falsely opened the door to spaCy PERSON mis-tags on
    # nearby place names. Kinship markers and "signatory" are unambiguous.
    r"\b(s/o|d/o|w/o|c/o|authorised signatory|authorized signatory)\b",
    re.IGNORECASE,
)


def _paragraph_group_node(nid, node_type, label, path_str, paragraphs, extra_meta=None):
    text = "\n".join(p.text for p in paragraphs)
    meta = {"atomic": len(paragraphs) == 1 and node_type != "paragraph"}
    if extra_meta:
        meta.update(extra_meta)
    if _SIGNATURE_ZONE_RE.search(text):
        meta["zone"] = "signature_block"
    return Node(nid, "docx", node_type, label, path_str, text,
                source_ref=paragraphs, meta=meta)


def parse_docx(path):
    doc = docx.Document(path)
    counter = [0]

    def next_id():
        counter[0] += 1
        return f"docx:{counter[0]}"

    root = Node(next_id(), "docx", "document", None, "", None)

    body = Node(next_id(), "docx", "body", None, "body", None)
    root.children.append(body)
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        node = _paragraph_group_node(next_id(), "paragraph", None, f"body.paragraph[{i}]", [p])
        body.children.append(node)
        for label, vs, ve, vtext in extract_pseudo_fields(p.text):
            node.children.append(Node(next_id(), "docx", "pseudo_field", label,
                                       f"{node.path}.field[{label}]", vtext,
                                       source_ref=[p], meta={"atomic": True, "span_in_parent": (vs, ve)}))

    tables_node = Node(next_id(), "docx", "tables", None, "tables", None)
    root.children.append(tables_node)
    for ti, table in enumerate(doc.tables):
        t_node = Node(next_id(), "docx", "table", None, f"table[{ti}]", None)
        tables_node.children.append(t_node)
        two_col = len(table.columns) == 2
        header_cells = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        expect_new_header = False
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue  # header row (or the two_col table's label baseline) isn't a data row

            cell_texts = [c.text.strip() for c in row.cells]
            distinct_non_blank = {t for t in cell_texts if t}
            # a real RHP table often bundles several logical mini-tables together,
            # separated by a merged full-width row acting as a section divider —
            # python-docx reports that merge as every cell sharing identical text
            is_divider = len(cell_texts) > 2 and len(distinct_non_blank) <= 1
            if is_divider:
                header_cells = None
                expect_new_header = True
                continue
            if expect_new_header and not two_col:
                header_cells = cell_texts
                expect_new_header = False
                continue

            row_node = Node(next_id(), "docx", "table_row", None, f"table[{ti}].row[{ri}]", None)
            t_node.children.append(row_node)
            # 2-column tables are often "Label | Value" per row (not header+records),
            # so use the row's own first cell as the label instead of a shared header
            row_label = row.cells[0].text.strip() if two_col and len(row.cells) == 2 else None
            for ci, cell in enumerate(row.cells):
                if two_col and ci == 0:
                    continue  # this cell IS the label, not a value to classify
                if two_col:
                    col_label = row_label
                else:
                    col_label = header_cells[ci] if header_cells and ci < len(header_cells) else None
                cell_paragraphs = cell.paragraphs or []
                cell_node = _paragraph_group_node(
                    next_id(), "table_cell", col_label,
                    f"table[{ti}].row[{ri}].cell[{ci}]", cell_paragraphs,
                    extra_meta={"col_index": ci},
                )
                row_node.children.append(cell_node)

    headers_node = Node(next_id(), "docx", "headers", None, "headers", None)
    footers_node = Node(next_id(), "docx", "footers", None, "footers", None)
    root.children.append(headers_node)
    root.children.append(footers_node)
    for si, section in enumerate(doc.sections):
        if any(p.text.strip() for p in section.header.paragraphs):
            headers_node.children.append(
                _paragraph_group_node(next_id(), "header", None, f"section[{si}].header",
                                       section.header.paragraphs))
        if any(p.text.strip() for p in section.footer.paragraphs):
            footers_node.children.append(
                _paragraph_group_node(next_id(), "footer", None, f"section[{si}].footer",
                                       section.footer.paragraphs))

    return root, doc


def _run_spans(paragraph):
    spans, pos = [], 0
    for run in paragraph.runs:
        spans.append((run, pos, pos + len(run.text)))
        pos += len(run.text)
    return spans


def _replace_span_in_paragraph(paragraph, start, end, new_text):
    touched = [t for t in _run_spans(paragraph) if t[2] > start and t[1] < end]
    if not touched:
        return
    first_run, first_start, _ = touched[0]
    last_run, last_start, last_end = touched[-1]
    local_start = start - first_start
    local_end = end - last_start
    if first_run is last_run:
        first_run.text = first_run.text[:local_start] + new_text + first_run.text[local_end:]
    else:
        first_run.text = first_run.text[:local_start] + new_text
        last_run.text = last_run.text[local_end:]
        for r, _, _ in touched[1:-1]:
            r.text = ""


def _apply_spans_to_paragraphs(paragraphs, joined_text, spans):
    """spans: [(start, end, replacement)] offsets into `joined_text`."""
    # locate paragraph boundaries within the joined ("\n"-separated) text
    bounds = []
    pos = 0
    for p in paragraphs:
        bounds.append((pos, pos + len(p.text)))
        pos += len(p.text) + 1  # +1 for the joining "\n"

    for start, end, repl in sorted(spans, key=lambda s: s[0], reverse=True):
        for p, (b_start, b_end) in zip(paragraphs, bounds):
            if start >= b_start and end <= b_end:
                _replace_span_in_paragraph(p, start - b_start, end - b_start, repl)
                break
            if start >= b_start and start < b_end:
                # span spilled past this paragraph's end (rare); clip to paragraph
                _replace_span_in_paragraph(p, start - b_start, b_end - b_start, repl)
                break


def write_docx(root, doc, output_path):
    # pseudo_field nodes never carry their own "redactions": the walker folds
    # their spans into the parent paragraph's coordinate space (see
    # redact.py's _detect_with_pseudo_fields), so only paragraph/cell/header/
    # footer nodes ever need writing here.
    for node in root.walk():
        spans = node.meta.get("redactions")
        if not spans or node.value is None or node.source_ref is None:
            continue
        _apply_spans_to_paragraphs(node.source_ref, node.value, spans)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _cluster_lines(words, tol=3):
    """Group words sharing similar 'top' into lines, left-to-right."""
    lines = []
    for w in sorted(words, key=lambda w: (round(w["top"] / tol), w["x0"])):
        placed = False
        for line in lines:
            if abs(line["top"] - w["top"]) <= tol:
                line["words"].append(w)
                line["top"] = min(line["top"], w["top"])
                line["bottom"] = max(line["bottom"], w["bottom"])
                placed = True
                break
        if not placed:
            lines.append({"top": w["top"], "bottom": w["bottom"], "words": [w]})
    for line in lines:
        line["words"].sort(key=lambda w: w["x0"])
    lines.sort(key=lambda l: l["top"])
    return lines


def _line_text_and_word_offsets(words):
    text_parts = []
    offsets = []  # (word, char_start, char_end)
    pos = 0
    for i, w in enumerate(words):
        if i > 0:
            text_parts.append(" ")
            pos += 1
        wt = w["text"]
        offsets.append((w, pos, pos + len(wt)))
        text_parts.append(wt)
        pos += len(wt)
    return "".join(text_parts), offsets


def parse_pdf(path):
    counter = [0]

    def next_id():
        counter[0] += 1
        return f"pdf:{counter[0]}"

    root = Node(next_id(), "pdf", "document", None, "", None)

    with pdfplumber.open(path) as pdf:
        for pi, page in enumerate(pdf.pages):
            page_node = Node(next_id(), "pdf", "page", None, f"page[{pi}]", None,
                              meta={"page_index": pi, "page_bbox": (0, 0, page.width, page.height)})
            root.children.append(page_node)

            tables = page.find_tables()
            table_bboxes = [t.bbox for t in tables]

            def in_any_table(word):
                cx = (word["x0"] + word["x1"]) / 2
                cy = (word["top"] + word["bottom"]) / 2
                return any(bx0 <= cx <= bx1 and by0 <= cy <= by1 for (bx0, by0, bx1, by1) in table_bboxes)

            words = page.extract_words()
            loose_words = [w for w in words if not in_any_table(w)]
            lines = _cluster_lines(loose_words)

            for li, line in enumerate(lines):
                text, offsets = _line_text_and_word_offsets(line["words"])
                if not text.strip():
                    continue
                line_bbox = (
                    min(w["x0"] for w in line["words"]),
                    line["top"],
                    max(w["x1"] for w in line["words"]),
                    line["bottom"],
                )
                line_node = Node(next_id(), "pdf", "line", None, f"page[{pi}].line[{li}]", text,
                                  source_ref={"page": pi},
                                  meta={"word_offsets": offsets, "bbox": line_bbox, "atomic": False})
                page_node.children.append(line_node)
                for label, vs, ve, vtext in extract_pseudo_fields(text):
                    line_node.children.append(Node(
                        next_id(), "pdf", "pseudo_field", label, f"{line_node.path}.field[{label}]", vtext,
                        source_ref={"page": pi}, meta={"atomic": True, "word_offsets": offsets,
                                                        "span_in_parent": (vs, ve), "bbox": line_bbox}))

            for ti, table in enumerate(tables):
                t_node = Node(next_id(), "pdf", "table", None, f"page[{pi}].table[{ti}]", None)
                page_node.children.append(t_node)
                grid = table.extract()
                header = [(c or "").strip() for c in grid[0]] if grid else []
                for ri, row in enumerate(table.rows):
                    if ri == 0:
                        continue
                    row_node = Node(next_id(), "pdf", "table_row", None,
                                     f"page[{pi}].table[{ti}].row[{ri}]", None)
                    t_node.children.append(row_node)
                    for ci, cell_bbox in enumerate(row.cells):
                        if cell_bbox is None:
                            continue
                        cx0, cy0, cx1, cy1 = cell_bbox
                        cell_words = [w for w in words if cx0 <= (w["x0"] + w["x1"]) / 2 <= cx1
                                      and cy0 <= (w["top"] + w["bottom"]) / 2 <= cy1]
                        text, offsets = _line_text_and_word_offsets(cell_words)
                        if not text.strip():
                            continue
                        col_label = header[ci] if ci < len(header) else None
                        row_node.children.append(Node(
                            next_id(), "pdf", "table_cell", col_label,
                            f"page[{pi}].table[{ti}].row[{ri}].cell[{ci}]", text,
                            source_ref={"page": pi},
                            meta={"atomic": True, "word_offsets": offsets, "bbox": cell_bbox}))

    return root, path


def write_pdf(root, input_path, output_path):
    # pseudo_field nodes never carry their own "redactions": the walker folds
    # their spans into the parent line's coordinate space (see redact.py's
    # _detect_with_pseudo_fields), so spans here are always already relative
    # to this node's own (line/table_cell) word_offsets.
    fitz_doc = fitz.open(input_path)
    for node in root.walk():
        spans = node.meta.get("redactions")
        if not spans or node.value is None:
            continue
        offsets = node.meta.get("word_offsets", [])
        page_index = node.source_ref["page"]
        page = fitz_doc[page_index]
        for start, end, repl in spans:
            hit_words = [w for (w, ws, we) in offsets if we > start and ws < end]
            if not hit_words:
                continue
            bbox = (
                min(w["x0"] for w in hit_words),
                min(w["top"] for w in hit_words),
                max(w["x1"] for w in hit_words),
                max(w["bottom"] for w in hit_words),
            )
            rect = fitz.Rect(bbox)
            page.add_redact_annot(rect, text=repl, fill=(1, 1, 1), fontsize=8)
    for page in fitz_doc:
        page.apply_redactions()
    fitz_doc.save(output_path)
    fitz_doc.close()


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def parse_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    counter = [0]

    def next_id():
        counter[0] += 1
        return f"txt:{counter[0]}"

    root = Node(next_id(), "txt", "document", None, "", None)

    for li, raw_line in enumerate(lines):
        text = raw_line.rstrip("\n")
        if not text.strip():
            continue
        m = _TXT_KV_RE.match(text) or _TXT_DASH_RE.match(text)
        if m:
            label, value = m.group(1).strip(), m.group(2)
            val_start = text.rfind(value)
            node = Node(next_id(), "txt", "pseudo_field", label, f"line[{li}]", value,
                        source_ref={"line": li}, meta={"atomic": True, "span_in_parent": (val_start, val_start + len(value))})
        else:
            node = Node(next_id(), "txt", "line", None, f"line[{li}]", text,
                        source_ref={"line": li}, meta={"atomic": False})
        root.children.append(node)

    return root, lines


def write_txt(root, lines, output_path):
    by_line = {}
    for node in root.walk():
        spans = node.meta.get("redactions")
        if not spans or node.value is None:
            continue
        li = node.source_ref["line"]
        base = node.meta.get("span_in_parent", (0, len(node.value)))[0]
        by_line.setdefault(li, []).extend((s + base, e + base, r) for s, e, r in spans)

    for li, spans in by_line.items():
        text = lines[li].rstrip("\n")
        newline = "\n" if lines[li].endswith("\n") else ""
        lines[li] = _splice(text, spans) + newline

    with open(output_path, "w", encoding="utf-8") as f:
        f.writelines(lines)
